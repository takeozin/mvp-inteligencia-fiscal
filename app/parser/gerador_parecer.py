"""
Gerador de Minuta de Parecer Jurídico — Tema 69 STF
Utiliza python-docx para montar a minuta de recuperação tributária.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _moeda(valor) -> str:
    """Formata float/Decimal como R$ 1.234,56"""
    try:
        v = float(valor or 0)
    except (TypeError, ValueError):
        v = 0.0
    return f'R$ {v:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')


def _pct(valor) -> str:
    try:
        return f'{float(valor or 0):.4f}%'
    except (TypeError, ValueError):
        return '0,0000%'


def gerar_parecer_docx(sped, parser, calculo, pasta_saida: str) -> str:
    """
    Gera a Minuta de Parecer Jurídico em DOCX e retorna o caminho completo.
    """
    os.makedirs(pasta_saida, exist_ok=True)
    nome_arquivo = (
        f'parecer_tema69_{sped.empresa_id}_{sped.id}_'
        f'{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    )
    caminho = os.path.join(pasta_saida, nome_arquivo)

    document = Document()

    # Configuração da página (Margens)
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Estilos padrão
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Função auxiliar para títulos
    def add_heading(text, level=1):
        h = document.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.name = 'Arial'
            if level == 1:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(27, 94, 32) # Verde escuro
            elif level == 2:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    # ── Cabeçalho ────────────────────────────────────────────────────────────
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MINUTA DE PARECER TRIBUTÁRIO\n')
    run.bold = True
    run.font.size = Pt(16)
    
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run('ANÁLISE DE RECUPERAÇÃO DE CRÉDITOS — TEMA 69 DO STF')
    run_sub.bold = True

    document.add_paragraph('\n')

    # ── 1. Qualificação ──────────────────────────────────────────────────────
    empresa_nome = parser.nome_empresa or sped.empresa.nome
    empresa_cnpj = parser.cnpj_empresa or sped.empresa.cnpj

    add_heading('1. QUALIFICAÇÃO DA CONSULENTE', level=2)
    p_qual = document.add_paragraph()
    p_qual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_qual.add_run('Consulente: ').bold = True
    p_qual.add_run(f'{empresa_nome}, pessoa jurídica de direito privado, inscrita no CNPJ sob o n.º {empresa_cnpj}.\n')
    p_qual.add_run('Objeto da Análise: ').bold = True
    p_qual.add_run('Avaliação da viabilidade e do quantum debeatur referente à exclusão do ICMS da base de cálculo do PIS e da COFINS, conforme entendimento do Supremo Tribunal Federal (STF).')

    document.add_paragraph()

    # ── 2. Breve Relato ──────────────────────────────────────────────────────
    add_heading('2. BREVE RELATO', level=2)
    p_relato = document.add_paragraph()
    p_relato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto_relato = (
        f'A presente análise baseia-se nas informações constantes nos arquivos SPED EFD-Contribuições '
        f'fornecidos pela Consulente (referência: {sped.nome_arquivo}), referentes ao '
        f'período de {parser.periodo_ini} a {parser.periodo_fim}. A Consulente submete-se à '
        f'tributação no regime {calculo.regime_calculo or "não especificado"}.'
    )
    p_relato.add_run(texto_relato)

    document.add_paragraph()

    # ── 3. Fundamentação Legal ───────────────────────────────────────────────
    add_heading('3. FUNDAMENTAÇÃO JURÍDICA', level=2)
    p_fund = document.add_paragraph()
    p_fund.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto_fund = (
        'O Supremo Tribunal Federal, ao julgar o Recurso Extraordinário (RE) n.º 574.706, submetido ao '
        'regime de repercussão geral (Tema 69), firmou a tese de que "O ICMS não compõe a base de cálculo '
        'para a incidência do PIS e da COFINS". Tal decisão transitou em julgado em 15/03/2017.\n\n'
        'Em sede de Embargos de Declaração, o STF (Ação Declaratória de Constitucionalidade - ADC 18) '
        'modulou os efeitos da decisão, estabelecendo que a tese produz efeitos a partir de 15/03/2017 '
        '(data do julgamento de mérito), ressalvadas as ações judiciais e procedimentos administrativos '
        'protocolados até essa data.\n\n'
        'Além disso, definiu-se que o ICMS a ser excluído da base de cálculo das referidas contribuições '
        'é o ICMS destacado nas notas fiscais, conforme a legislação vigente e a jurisprudência pacificada.'
    )
    p_fund.add_run(texto_fund)

    document.add_paragraph()

    # ── 4. Apuração dos Valores ──────────────────────────────────────────────
    add_heading('4. APURAÇÃO PRELIMINAR DE CRÉDITOS', level=2)
    p_apura = document.add_paragraph()
    p_apura.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_apura.add_run(
        'A partir da extração dos dados dos documentos fiscais, segregamos o ICMS destacado nas saídas '
        'e recalculamos a base do PIS e da COFINS, subtraindo o respectivo valor. O quadro resumo abaixo '
        'demonstra a diferença entre o valor apurado originalmente e o valor devido:'
    )

    # Tabela de Valores
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rubrica'
    hdr_cells[1].text = 'PIS'
    hdr_cells[2].text = 'COFINS'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    registros_tabela = [
        ('Base Declarada', _moeda(calculo.vl_bc_pis_declarado), _moeda(calculo.vl_bc_cofins_declarado)),
        ('ICMS Destacado Excluído', _moeda(float(calculo.vl_bc_pis_declarado or 0) - float(calculo.vl_bc_pis_corrigido or 0)), _moeda(float(calculo.vl_bc_cofins_declarado or 0) - float(calculo.vl_bc_cofins_corrigido or 0))),
        ('Base Corrigida', _moeda(calculo.vl_bc_pis_corrigido), _moeda(calculo.vl_bc_cofins_corrigido)),
        ('Tributo Devido', _moeda(calculo.vl_pis_corrigido), _moeda(calculo.vl_cofins_corrigido)),
        ('Tributo Declarado', _moeda(calculo.vl_pis_declarado), _moeda(calculo.vl_cofins_declarado)),
        ('CRÉDITO APURADO', _moeda(calculo.credito_pis), _moeda(calculo.credito_cofins)),
    ]

    for rubrica, pis, cofins in registros_tabela:
        row_cells = table.add_row().cells
        row_cells[0].text = rubrica
        row_cells[1].text = pis
        row_cells[2].text = cofins

    document.add_paragraph()
    
    p_total = document.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_total = p_total.add_run(f'TOTAL ESTIMADO DO CRÉDITO: {_moeda(calculo.credito_total)}')
    run_total.bold = True
    run_total.font.size = Pt(14)
    run_total.font.color.rgb = RGBColor(27, 94, 32)

    document.add_paragraph()

    # ── 5. Conclusão ─────────────────────────────────────────────────────────
    add_heading('5. CONCLUSÃO', level=2)
    p_conclusao = document.add_paragraph()
    p_conclusao.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto_conclusao = (
        'Ante o exposto, conclui-se que a Consulente possui viabilidade jurídica para pleitear a '
        'recuperação e compensação dos valores recolhidos a maior, na ordem de '
        f'{_moeda(calculo.credito_total)}, correspondentes ao período analisado.\n\n'
        'Recomenda-se a retificação das obrigações acessórias (EFD-Contribuições e DCTF) ou o ingresso '
        'na via administrativa via PER/DCOMP, a depender da estratégia adotada pela diretoria da Consulente. '
        'Os valores ora apresentados estão sujeitos à atualização pela Taxa SELIC e devem passar por '
        'auditoria final antes do aproveitamento.'
    )
    p_conclusao.add_run(texto_conclusao)

    document.add_paragraph('\n\n\n')

    # ── Assinatura ───────────────────────────────────────────────────────────
    p_ass = document.add_paragraph()
    p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ass.add_run('_________________________________________________\n')
    p_ass.add_run('DEPARTAMENTO JURÍDICO / TRIBUTÁRIO\n')
    p_ass.add_run(f'{datetime.now().strftime("%d de %B de %Y")}')

    # Salva o arquivo
    document.save(caminho)
    return caminho
